# ============================================================
# MILESTONE 3: AUTOMATED DRAFT GENERATION (FINAL – STABLE)
# Results → Gemini ONLY (Runtime Model Selection)
# ============================================================

import os
import json
from typing import List, Dict
from datetime import datetime
from dotenv import load_dotenv

# ---------------- OPTIONAL EXPORT DEPENDENCIES ----------------
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.platypus import SimpleDocTemplate, Paragraph
    from reportlab.lib.styles import getSampleStyleSheet
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

load_dotenv()

METADATA_DIR = "processed/metadata"
OUTPUT_DIR = "milestone3_output"
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ============================================================
# 1. LOAD METADATA
# ============================================================

def load_all_paper_metadata() -> List[Dict]:
    papers = []
    if not os.path.exists(METADATA_DIR):
        raise RuntimeError("❌ processed/metadata not found")

    for f in os.listdir(METADATA_DIR):
        if f.endswith(".json"):
            with open(os.path.join(METADATA_DIR, f), "r", encoding="utf-8") as file:
                papers.append(json.load(file))

    print(f"📄 Loaded {len(papers)} processed papers")
    return papers

# ============================================================
# 2. SYNTHESIS
# ============================================================

def synthesize_findings(papers: List[Dict]) -> Dict[str, List[str]]:
    synthesis = {}
    for p in papers:
        for theme, findings in p.get("themes", {}).items():
            synthesis.setdefault(theme, []).extend(findings)

    for k in synthesis:
        synthesis[k] = list(set(synthesis[k]))
    return synthesis

# ============================================================
# 3. CITATIONS
# ============================================================

def build_citation_key(paper: Dict) -> str:
    authors = paper.get("authors", [])
    year = paper.get("year", "n.d.")
    last = authors[0].split()[-1] if authors else "Unknown"
    return f"{last} et al., {year}"

def build_citation_map(papers: List[Dict]) -> Dict[int, str]:
    return {i: build_citation_key(p) for i, p in enumerate(papers)}

# ============================================================
# 4. PROMPTS
# ============================================================

def build_prompt(section: str, content: str) -> str:
    return f"""
Write a formal academic {section} section.

{content}
"""

# ============================================================
# 5. GEMINI HELPERS
# ============================================================

def auto_select_gemini_model(client):
    models = client.models.list()
    for m in models:
        try:
            client.models.generate_content(
                model=m.name,
                contents="Test"
            )
            return m.name
        except Exception:
            continue
    raise RuntimeError("❌ No usable Gemini model found")

def generate_with_gemini(prompt: str, api_key: str) -> str:
    from google import genai
    client = genai.Client(api_key=api_key)
    model = auto_select_gemini_model(client)

    response = client.models.generate_content(
        model=model,
        contents=prompt
    )
    return response.text.strip()

def generate_results_with_gemini_only(prompt: str) -> str:
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        raise RuntimeError("❌ GEMINI_API_KEY required")

    from google import genai
    client = genai.Client(api_key=api_key)
    model = auto_select_gemini_model(client)

    print(f"✅ Using Gemini model: {model}")

    response = client.models.generate_content(
        model=model,
        contents=prompt
    )
    return response.text.strip()

# ============================================================
# 6. FALLBACK GENERATOR
# ============================================================

def generate_with_llm(prompt: str) -> str:
    key = os.getenv("GEMINI_API_KEY")
    if key:
        try:
            return generate_with_gemini(prompt, key)
        except Exception:
            print("⚠ Gemini unavailable, using fallback.")

    return "This section was generated using a fallback mechanism."

# ============================================================
# 7. SECTIONS
# ============================================================

def generate_abstract(synthesis: Dict) -> str:
    return generate_with_llm(build_prompt("abstract", str(synthesis)))

def generate_methods(papers: List[Dict]) -> str:
    methods = [
        p.get("sections", {}).get("methods", "")
        for p in papers if p.get("sections", {}).get("methods")
    ]
    return generate_with_llm(build_prompt("methods", "\n".join(methods)))

def generate_results(synthesis: Dict, citation_map: Dict) -> str:
    cited = []
    for i, (theme, findings) in enumerate(synthesis.items()):
        if findings:
            cited.append(f"{findings[0]} ({citation_map[i % len(citation_map)]})")

    return generate_results_with_gemini_only(
        build_prompt("results", "\n".join(cited))
    )

# ============================================================
# 8. REFERENCES
# ============================================================

def generate_references(papers: List[Dict]) -> List[str]:
    refs = []
    for p in papers:
        authors = ", ".join(p.get("authors", [])[:3])
        refs.append(
            f"{authors} ({p.get('year','n.d.')}). {p.get('title','')}."
        )
    return refs

# ============================================================
# 9. EXPORTS
# ============================================================

def export_to_docx(draft: Dict):
    if not DOCX_AVAILABLE:
        return
    doc = Document()
    doc.add_heading("Automated Research Draft", 0)

    for sec in ["abstract", "methods", "results"]:
        doc.add_heading(sec.capitalize(), 1)
        doc.add_paragraph(draft[sec])

    doc.add_heading("References", 1)
    for r in draft["references"]:
        doc.add_paragraph(r, style="List Number")

    doc.save(os.path.join(OUTPUT_DIR, "final_draft.docx"))

def export_to_pdf(draft: Dict):
    if not PDF_AVAILABLE:
        return

    styles = getSampleStyleSheet()
    story = []

    for sec in ["abstract", "methods", "results"]:
        story.append(Paragraph(f"<b>{sec.capitalize()}</b>", styles["Heading2"]))
        story.append(Paragraph(draft[sec], styles["Normal"]))

    SimpleDocTemplate(
        os.path.join(OUTPUT_DIR, "final_draft.pdf")
    ).build(story)

# ============================================================
# 10. MAIN
# ============================================================

def run_milestone_3():
    print("\n🚀 Starting Milestone 3\n")

    papers = load_all_paper_metadata()
    synthesis = synthesize_findings(papers)
    citation_map = build_citation_map(papers)

    output = {
        "generated_at": datetime.now().isoformat(),
        "abstract": generate_abstract(synthesis),
        "methods": generate_methods(papers),
        "results": generate_results(synthesis, citation_map),
        "references": generate_references(papers)
    }

    with open(os.path.join(OUTPUT_DIR, "final_draft.json"), "w", encoding="utf-8") as f:
        json.dump(output, f, indent=4, ensure_ascii=False)

    export_to_docx(output)
    export_to_pdf(output)

    print("✅ Milestone 3 completed successfully!")

if __name__ == "__main__":
    run_milestone_3()
